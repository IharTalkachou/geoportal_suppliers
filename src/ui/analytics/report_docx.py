import streamlit as st
import pandas as pd
import io
import json
from datetime import datetime, date, timedelta
from sqlalchemy import text
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from config.cache import query_db, clear_cache

MONTHS_RU = {
    1: "Январь", 2: "Февраль", 3: "Март", 4: "Апрель", 5: "Май", 6: "Июнь",
    7: "Июль", 8: "Август", 9: "Сентябрь", 10: "Октябрь", 11: "Ноябрь", 12: "Декабрь"
}

# 🟢 Новая структура с принудительным порядком (ключи с префиксами)
DEFAULT_SECTIONS = {
    "6.1": {
        "title": "6.1. Отчёт о работах по ведению (эксплуатации) Национального геопортала.",
        "blocks": {
            "01_611_header": {
                "title": "Заголовок 6.1.1 (Администрирование)",
                "active": True,
                "content": "6.1.1. Администрирование Национального геопортала, в том числе:\nприем, рассмотрение заявок на регистрацию и регистрация пользователей на Национальном геопортале:"
            },
            "02_reg_users": {
                "title": "Блок: Регистрация (пользователи)",
                "active": True,
                "content": ""
            },
            "03_cabinets": {
                "title": "Блок: Электронные кабинеты",
                "active": True,
                "content": ""
            },
            "04_total_stats": {
                "title": "Блок: Статистика на конец периода",
                "active": True,
                "content": ""
            },
            "05_provision_nipd": { 
                "title": "Блок: Предоставление НИПД (заявки)", 
                "active": True, 
                "content": "прием, рассмотрение заявок на предоставление в пользование наборов пространственных данных, включенных в НИПД:" 
            },
            "06_metadata_nipd": {
                "title": "Блок: Рецензирование метаданных (НИПД) | записей в дневнике Насти + ES changeDate + ES createDate + другие (если есть что - указаны в 6.1.3)",
                "active": True,
                "content": "рецензирование (проверка) и публикация на Национальном геопортале метаданных о наборах пространственных данных, включенных в НИПД, и сервисов для этих наборов – за отчетный период в Каталоге метаданных размещено [X] записей;"
            },
            "07_metadata_gkf": {
                "title": "Блок: Метаданные (ГКГФ) | Администрирование - Отчёты",
                "active": True,
                "content": "размещение и публикация на Национальном геопортале метаданных о материалах и данных Госкартгеофонда – опубликовано (обновлено) [X] записей метаданных;"
            },
            "08_moderation": {
                "title": "Блок: Модерация (консультации и информирование)",
                "active": True,
                "content": "модерация Национального геопортала (работа с комментариями и сообщениями пользователей), в том числе:\n– обработка сообщений обратной связи от пользователей по обнаруженным ошибкам или предложениям по улучшению работы Национального геопортала – всего поступило [X] сообщений от ([Наименование пользователей]);\n– консультирование Поставщиков ([Список]) о способах и порядке создания записей метаданных в каталоге метаданных, заполнении разделов метаданных, рабочих процессах каталога метаданных и порядке валидации записей;\n– информирование пользователей ([Список]) о Национальной инфраструктуре пространственных данных, её типовых компонентах, порядке подачи заявок на регистрацию на Национальном геопортале и перспективных направлениях сотрудничества и обмена информацией;"
            },
            "09_612_header": {
                "title": "Заголовок 6.1.2", 
                "active": True, 
                "content": "6.1.2. Заключение соглашений о доступе, обмене и использовании наборов пространственных данных между оператором Национального геопортала и поставщиками наборов пространственных данных, включенных в НИПД, в том числе:"
            },
            "10_meetings": {
                "title": "Блок 10: Совещания", 
                "active": True, 
                "content": ""
            },
            "11_approval_sent": {
                "title": "Блок 11: Направлено на согласование", 
                "active": True, 
                "content": ""
            },
            "12_agreement_signed": {
                "title": "Блок 12: Подписано Соглашение", 
                "active": True, 
                "content": ""
            },
            "13_protocol_signed": {
                "title": "Блок 13: Подписан Протокол", 
                "active": True, 
                "content": ""
            },
            "14_letters_gov": {
                "title": "Блок 14: Письма (Гос.органы)", 
                "active": True, 
                "content": ""
            },
            "15_letters_onpd": {
                "title": "Блок 15: Письма (ОНПД)", 
                "active": True, 
                "content": ""
            },
            "16_letters_other": {
                "title": "Блок 16: Письма (Прочие)", 
                "active": True, 
                "content": ""
            },
            "17_website_content": {
                "title": "Блок 17: Содержание сайта и новости",
                "active": True,
                "content": (
                    "6.1.3. Подготовка, публикация, обновление (редактирование) содержания сайта "
                    "Национального геопортала, относящегося к функционированию НИПД, в том числе:\n"
                    "– подготовлены и опубликованы на сайте и в Телеграм-канале Национального геопортала "
                    "новости о заключении Соглашений с поставщиками наборов пространственных данных;\n"
                    "– созданы и опубликованы метаданные о … – всего [Х] записей: …\n"
                    "– созданы и опубликованы инструкции о …;"
                )
            },
            "18_support": {
                "title": "Блок 18: Методическая поддержка",
                "active": True,
                "content": (
                    "6.1.4. Методическая и консультативная поддержка поставщиков и пользователей, в том числе:\n"
                    "– внесены исправления в документацию Каталога метаданных Национального геопортала в части …, "
                    "новая версия документации размещена в тестовой среде Национального геопортала;\n"
                    "– проведено обучение специалистов Поставщика ([Список]) по редактированию данных "
                    "в среде Национального геопортала;\n"
                    "– выполнена настройка среды редактирования пространственных данных в Каталоге геоданных "
                    "Национального геопортала для Поставщика ([Список]);\n"
                    "– осуществлена техническая поддержка пользователей Национального геопортала – всего поступило "
                    "[Х] обращений:\n[Дата] – [Лицо] – [Вопрос] – [Комментарий]"
                )
            }
        }
    },
    "6.3": {
        "title": "6.3. Отчёт о работах по анализу функционирования НИПД.",
        "blocks": {
            "19_stats_631": {
                "title": "Блок 19: Посещаемость (6.3.1)",
                "active": True,
                "is_structured": True, # 👈 Флаг для полей ввода
                "content": ""
            },
            "20_popularity_632": {
                "title": "Блок 20: Востребованность данных (6.3.2)",
                "active": True,
                "is_structured": True, # 👈 Флаг для полей ввода
                "content": ""
            }
        }
    }
}

def pluralize(n, forms):
    """Склонение существительных: [заявка, заявки, заявок]"""
    n = abs(int(n))
    if n % 10 == 1 and n % 100 != 11: return forms[0]
    elif 2 <= n % 10 <= 4 and (n % 100 < 10 or n % 100 >= 20): return forms[1]
    else: return forms[2]

def pluralize_verb_fem(n, forms):
    """Склонение глаголов (была/были, выполнена/выполнены)"""
    n = abs(int(n))
    return forms[0] if (n % 10 == 1 and n % 100 != 11) else forms[1]

def pluralize_outcome(n, variant):
    """Склонение категорий исходов: выполнена/выполнены и т.д."""
    n = abs(int(n))
    forms = {
        "done": ["выполнена", "выполнены"],
        "op": ["отклонена Оператором", "отклонены Оператором"],
        "sup": ["отказана Поставщиком", "отказаны Поставщиками"]
    }
    return forms[variant][0] if (n % 10 == 1 and n % 100 != 11) else forms[variant][1]

def get_supplier_header(count, context_key):
    """
    Возвращает правильное окончание для слова Поставщик в зависимости от контекста.
    context_key: 
      'with' -> 'с Поставщиком' / 'с Поставщиками'
      'to'   -> 'Поставщику' / 'Поставщикам'
      'for'  -> 'Поставщику' / 'Поставщикам' (на согласование кому?)
    """
    if count <= 1:
        headers = {
            'with': 'с Поставщиком пространственных данных',
            'to': 'Поставщику обязательного набора пространственных данных',
            'for': 'Поставщику обязательных наборов пространственных данных',
            'for_general': 'Поставщику пространственных данных'
        }
    else:
        headers = {
            'with': 'с Поставщиками пространственных данных',
            'to': 'Поставщикам обязательных наборов пространственных данных',
            'for': 'Поставщикам обязательных наборов пространственных данных',
            'for_general': 'Поставщикам наборов пространственных данных'
        }
    return headers.get(context_key, 'Поставщикам')

def get_gkf_material_names(material_ids):
    """Превращает [1, 2] в 'Аэрофотоснимки, Ортофотопланы'"""
    if not material_ids: return "не указаны"
    # Принудительно приводим к списку целых чисел, если пришел JSON
    if isinstance(material_ids, str):
        import json
        try: ids = json.loads(material_ids)
        except: return "ошибка формата"
    else: ids = material_ids
    
    if not ids: return "не указаны"
    
    # Запрашиваем названия материалов
    placeholders = ', '.join([':id' + str(i) for i in range(len(ids))])
    params = {f'id{i}': int(v) for i, v in enumerate(ids)}
    res = query_db(f"SELECT material_name FROM ref_gkf_materials WHERE material_id IN ({placeholders})", params)
    
    return ", ".join(res['material_name'].tolist()) if not res.empty else "не найдены"

@st.dialog("Фиксация отчёта")
def confirm_fix_dialog(session, report_id, updated_full_data):
    st.write("Укажите дату и время фиксации отчёта. С этого момента начнётся сбор данных для отчёта за следующий месяц.")
    fix_date = st.date_input("Дата фиксации", value=date.today(), key="fix_report_date")
    fix_time = st.time_input("Время фиксации", value=datetime.now().time(), key="fix_report_time")

    if st.button("🔒 Подтвердить фиксацию", type="primary", width='stretch'):
        fixed_at = datetime.combine(fix_date, fix_time)
        try:
            session.execute(
                text("UPDATE reports_monthly SET fixed_at=:f, sections_data=:d WHERE report_id=:id"),
                {"f": fixed_at, "d": json.dumps(updated_full_data, ensure_ascii=False), "id": report_id}
            )
            session.commit()
            clear_cache()
            st.rerun()
        except Exception as e:
            st.error(f"Ошибка фиксации: {e}")
            session.rollback()

def get_report_period_bounds(report_month_date):
    prev = query_db("""
        SELECT fixed_at FROM reports_monthly 
        WHERE report_month < :d AND fixed_at IS NOT NULL 
        ORDER BY report_month DESC LIMIT 1
    """, {"d": report_month_date})
    start_time = prev.iloc[0]['fixed_at'] if not prev.empty else datetime.combine(report_month_date, datetime.min.time())
    return start_time, datetime.now()

def save_report_metric(session, report_id, key, value=None, text_val=None):
    """Сохраняет или обновляет метрику в БД"""
    session.execute(text("""
        INSERT INTO report_metrics (report_id, metric_key, metric_value, metric_text)
        VALUES (:rid, :key, :v, :t)
        ON CONFLICT (report_id, metric_key) DO UPDATE 
        SET metric_value = EXCLUDED.metric_value, metric_text = EXCLUDED.metric_text
    """), {"rid": report_id, "key": key, "v": value, "t": text_val})

def load_report_metrics(report_id):
    """Загружает все метрики отчета в словарь"""
    df = query_db("SELECT metric_key, metric_value, metric_text FROM report_metrics WHERE report_id = :id", {"id": report_id})
    if df.empty: return {}
    res = {}
    for _, r in df.iterrows():
        res[r['metric_key']] = r['metric_text'] if r['metric_text'] is not None else float(r['metric_value'])
    return res

def fetch_registration_stats(start_t, end_t):
    new_users = query_db("""
        SELECT r.applicant_type, r.applicant_name, r.org_type,
               (SELECT COUNT(*) FROM reg_request_users WHERE req_id = r.req_id) as u_count
        FROM reg_requests r
        WHERE r.processed_at BETWEEN :s AND :e AND r.status = 'Завершена'
    """, {"s": start_t, "e": end_t})
    total_stats = query_db("""
        SELECT 
            COUNT(*) FILTER (WHERE applicant_type = 'Физическое лицо') as phys,
            COUNT(*) FILTER (WHERE applicant_type = 'Юридическое лицо') as orgs
        FROM reg_requests WHERE processed_at <= :e AND status = 'Завершена'
    """, {"e": end_t}).iloc[0]
    return new_users, total_stats

def generate_docx_file(report_date, sections_data, month_name):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(2), Cm(2)
        section.left_margin, section.right_margin = Cm(3), Cm(1.5)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f"Отчёт за {month_name} {report_date.year} г.")
    run.bold = True

    for g_key in sorted(sections_data.keys()):
        group = sections_data[g_key]
        active_blocks = [b for b in group['blocks'].values() if b['active'] and b['content'].strip()]
        if not active_blocks: continue

        # 🟢 Настройка заголовка группы (не жирный, с отступом)
        h = doc.add_paragraph()
        h_run = h.add_run(group['title'])
        h_run.bold = False 
        '''h.paragraph_format.first_line_indent = Cm(1.25)
        h.paragraph_format.line_spacing = 1.0'''
        fmt_h = h.paragraph_format
        fmt_h.first_line_indent = Cm(1.25)
        fmt_h.line_spacing = 1.0
        fmt_h.space_after = Pt(0) 
        fmt_h.space_before = Pt(0)

        for b_key in sorted(group['blocks'].keys()):
            block = group['blocks'][b_key]
            if not block['active'] or not block['content'].strip(): continue
            
            paragraphs = block['content'].split('\n')
            for p_text in paragraphs:
                if p_text.strip():
                    p = doc.add_paragraph()
                    fmt = p.paragraph_format
                    # 🟢 Межстрочный интервал 1.0
                    fmt.line_spacing = 1.0
                    fmt.space_after = Pt(0)
                    
                    if p_text.startswith('\t'):
                        # Реальная табуляция без маркированного списка
                        p.text = p_text.replace('\t', '', 1)
                        fmt.first_line_indent = Pt(36) # Имитация табуляции через отступ
                    elif p_text.strip().startswith(('–', '-', '*')):
                        p.style = 'List Bullet'
                        p.text = p_text.strip()[1:].strip()
                    else:
                        p.text = p_text
                        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        fmt.first_line_indent = Cm(1.25)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render_monthly_report_tab(session):
    st.subheader("📄 Конструктор отчёта НИПД")

    # --- 1. ВЫБОР ПЕРИОДА ---
    col_y, col_m = st.columns(2)
    with col_y: 
        s_year = st.selectbox("Год", [2025, 2026, 2027], index=1)
    with col_m: 
        s_month = st.selectbox("Месяц", list(MONTHS_RU.keys()), format_func=lambda x: MONTHS_RU[x], index=datetime.now().month - 1)

    report_date = date(s_year, s_month, 1)
    report_res = query_db("SELECT * FROM reports_monthly WHERE report_month = :d", {"d": report_date})

    # --- 2. ИНИЦИАЛИЗАЦИЯ ЧЕРНОВИКА ---
    if report_res.empty:
        if st.button("➕ Создать структуру черновика", width='stretch'):
            session.execute(text("INSERT INTO reports_monthly (report_month, sections_data, created_by) VALUES (:d, :data, :uid)"),
                            {"d": report_date, "data": json.dumps(DEFAULT_SECTIONS, ensure_ascii=False), "uid": st.session_state.auth["user_id"]})
            session.commit(); clear_cache(); st.rerun()
        return

    rep_data = report_res.iloc[0]
    report_id = int(rep_data['report_id'])
    sections = rep_data['sections_data']
    is_fixed = rep_data['fixed_at'] is not None
    start_p, end_p = get_report_period_bounds(report_date)

    # 🟢 НОВОЕ: Принудительная загрузка метрик из БД в Session State
    # Это гарантирует, что при смене месяца поля не будут пустыми
    current_metrics = load_report_metrics(report_id)
    prefix = f"m_{report_id}_"
    for m_key, m_val in current_metrics.items():
        state_key = prefix + m_key
        if state_key not in st.session_state:
            st.session_state[state_key] = m_val

    if is_fixed: 
        st.warning(f"🔒 Отчёт зафиксирован {rep_data['fixed_at'].strftime('%d.%m.%Y %H:%M')}")
    st.caption(f"Данные собираются за период: {start_p.strftime('%d.%m.%Y %H:%M')} — {end_p.strftime('%d.%m.%Y %H:%M')}")

    # --- 3. СИНХРОНИЗАЦИЯ СТРУКТУРЫ И ПЕРВИЧНАЯ ЗАГРУЗКА В STATE ---
    structure_changed = False
    for g_key, g_val in DEFAULT_SECTIONS.items():
        if g_key not in sections:
            sections[g_key] = g_val
            structure_changed = True
        else:
            # 🟢 НОВОЕ: Обновляем заголовок группы, если он изменился в коде
            if sections[g_key].get('title') != g_val['title']:
                sections[g_key]['title'] = g_val['title']
                structure_changed = True

        for b_key, b_val in g_val['blocks'].items():
            if b_key not in sections[g_key]['blocks']:
                sections[g_key]['blocks'][b_key] = b_val
                structure_changed = True
            else:
                # 🟢 НОВОЕ: Обновляем заголовок блока, если он изменился в коде
                if sections[g_key]['blocks'][b_key].get('title') != b_val['title']:
                    sections[g_key]['blocks'][b_key]['title'] = b_val['title']
                    structure_changed = True
            
            # Инициализируем Session State
            t_key = f"t_{g_key}_{b_key}"
            tx_key = f"tx_{g_key}_{b_key}_{report_id}"
            
            if t_key not in st.session_state:
                st.session_state[t_key] = sections[g_key]['blocks'][b_key]['active']
            if tx_key not in st.session_state:
                st.session_state[tx_key] = sections[g_key]['blocks'][b_key]['content']
    
    if structure_changed:
        session.execute(text("UPDATE reports_monthly SET sections_data = :d WHERE report_id = :id"),
                        {"d": json.dumps(sections, ensure_ascii=False), "id": report_id})
        session.commit()

    # --- 4. ЕДИНАЯ КНОПКА СБОРА ДАННЫХ ---
    if not is_fixed:
        if st.button("✨ Собрать данные для отчёта из базы", width='stretch', type="secondary"):
            # А. Получаем данные из базы
            new_users, totals = fetch_registration_stats(start_p, end_p)

            # --- БЛОК 02: РЕГИСТРАЦИЯ ---
            phys = len(new_users[new_users['applicant_type'] == 'Физическое лицо'])
            org_u = new_users[(new_users['applicant_type'] == 'Юридическое лицо') & (new_users['org_type'] == 'Пользователь')]
            if phys > 0 or not org_u.empty:
                p_txt = f"{phys}-х физических лиц" if phys > 1 else f"1-го физического лица" if phys == 1 else ""
                o_txt = f"{len(org_u)}-х юридических лиц" if len(org_u) > 1 else f"1-го юридического лица" if len(org_u) == 1 else ""
                conn = " и " if (p_txt and o_txt) else ""
                det = ", ".join([f"{r['applicant_name']} – {r['u_count']} учётных записей" for _, r in org_u.iterrows()])
                st.session_state[f"tx_6.1_02_reg_users_{report_id}"] = f"\t– обработаны заявки, осуществлена регистрация {p_txt}{conn}{o_txt}{' (' + det + ')' if det else ''};"
                st.session_state[f"t_6.1_02_reg_users"] = True
            else:
                st.session_state[f"t_6.1_02_reg_users"] = False

            # --- БЛОК 03: КАБИНЕТЫ ---
            sups = new_users[(new_users['applicant_type'] == 'Юридическое лицо') & (new_users['org_type'] == 'Поставщик')]
            if not sups.empty:
                names = ", ".join(sups['applicant_name'].tolist())
                st.session_state[f"tx_6.1_03_cabinets_{report_id}"] = f"– созданы и настроены электронные кабинеты, включая настройку ролей, пользователей и шаблонов метаданных для {len(sups)}{'-го Поставщика' if len(sups) == 1 else '-х Поставщиков'}: {names};"
                st.session_state[f"t_6.1_03_cabinets"] = True
            else:
                st.session_state[f"t_6.1_03_cabinets"] = False

            # --- БЛОК 04: СТАТИСТИКА ГП ---
            st.session_state[f"tx_6.1_04_total_stats_{report_id}"] = f"на конец отчётного периода на Национальном геопортале зарегистрировано:\n\t– {totals['phys']} физических лиц;\n\t– {totals['orgs']} юридических лиц;"
            st.session_state[f"t_6.1_04_total_stats"] = True

            # --- БЛОК 05: ПРЕДОСТАВЛЕНИЕ НИПД ---
            stats = fetch_provision_stats(start_p, end_p)

            def get_outcome_category(row):
                code = row['stage_code']
                comm = str(row['last_comment']).lower() if row['last_comment'] else ""
                # Если закрыта, но в комментарии есть слово "отказ" - это отказ
                if code == 'REQ_CLOSE' and ('отказ' in comm or 'возврат' in comm):
                    return "отказана Поставщиком"
                if code in ('REQ_AGREE_COMPL', 'REQ_CLOSE'): return "выполнена"
                if code == 'REQ_REGIS_RETUR' or code == 'REQ_REFUS_SUBMI': return "отклонена Оператором"
                return "отказана Поставщиком"

            def format_passport(row, is_nipd=True):
                applicant = row['applicant_name']
                if is_nipd:
                    details = f"заявитель — {applicant}, набор «{row['dataset_name']}» ({row['info_name']}), поставщик — {row['supplier_name']}"
                else:
                    mats = get_gkf_material_names(row['gkf_material_ids'])
                    details = f"заявитель — {applicant}, материалы: {mats}"
                
                cat = get_outcome_category(row)
                variant = "done" if "выполн" in cat else "op" if "Оператор" in cat else "sup"
                outcome_text = pluralize_outcome(1, variant)
                
                if variant != "done":
                    reason = f", причина: {row['last_comment']}" if row['last_comment'] else ""
                    return f"{outcome_text} ({details}{reason})"
                return f"{outcome_text} ({details})"

            def build_section(data, label, is_nipd=True):
                proc = data['processed']
                work = data['in_work']
                
                if proc.empty and work.empty:
                    return f"– {label} — в отчетном периоде не поступало;"

                count_proc = len(proc)
                v_obrab = pluralize_verb_fem(count_proc, ['полностью обработана', 'полностью обработаны'])
                n_curr = f"{count_proc} {pluralize(count_proc, ['заявка', 'заявки', 'заявок'])}"
                
                if count_proc >= 3:
                    bits = []
                    # Группируем для краткой формы
                    results = [get_outcome_category(r) for _, r in proc.iterrows()]
                    d_n = results.count("выполнена")
                    o_n = results.count("отклонена Оператором")
                    s_n = results.count("отказана Поставщиком")
                    if d_n: bits.append(f"{d_n} — {pluralize_outcome(d_n, 'done')}")
                    if o_n: bits.append(f"{o_n} — {pluralize_outcome(o_n, 'op')}")
                    if s_n: bits.append(f"{s_n} — {pluralize_outcome(s_n, 'sup')}")
                    res_str = f" ({', '.join(bits)})"
                else:
                    res_str = f" ({', '.join([format_passport(r, is_nipd) for _, r in proc.iterrows()])})"

                # Логика "В работе": Скрываем, если пусто
                work_str = ""
                if not work.empty:
                    w_list = []
                    for _, r in work.iterrows():
                        if is_nipd: w_list.append(f"заявитель — {r['applicant_name']}, «{r['dataset_name']}»")
                        else: w_list.append(f"заявитель — {r['applicant_name']}, материалы: {get_gkf_material_names(r['gkf_material_ids'])}")
                    
                    w_desc = f" ({'; '.join(w_list)})" if len(work) < 3 else ""
                    work_str = f", по {len(work)} {pluralize(len(work), ['заявке', 'заявкам', 'заявкам'])}{w_desc} работа продолжается"

                return f"– {label} — {v_obrab} {n_curr}{res_str}{work_str};"

            text_nipd = build_section(stats['NIPD'], "включенных в НИПД", True)
            text_gkf = build_section(stats['GKF'], "материалов Госкартгеофонда", False) # 👈 Одно слово "материалов"
            
            def format_total(count, label_text):
                if count == 0: return f"– {label_text} — не получены"
                return f"– {count} {pluralize(count, ['заявка', 'заявки', 'заявок'])} {label_text}"

            t_nipd = format_total(stats['NIPD']['totals']['total_received'], "на предоставление в пользование наборов пространственных данных, включенных в НИПД")
            t_gkf = format_total(stats['GKF']['totals']['total_received'], "на предоставление в пользование материалов Госкартгеофонда")

            full_content = (
                "прием, рассмотрение заявок на предоставление в пользование наборов пространственных данных:\n"
                f"{text_nipd}\n{text_gkf}\n"
                "на конец отчётного периода обработано:\n"
                f"{t_nipd},\n{t_gkf};"
            )

            st.session_state[f"tx_6.1_05_provision_nipd_{report_id}"] = full_content
            st.session_state[f"t_6.1_05_provision_nipd"] = True
            
            # --- БЛОКИ 06 - 08: ввод данных из ElasticSearch, модерация и техподдержка
            # нужно проработать

            # --- БЛОКИ 09 - 16: БЮРОКРАТИЯ 6.1.2 ---
            b_stats = fetch_bureaucracy_logic_stats(start_p, end_p)
            
            # Блок 10
            is_b10 = len(b_stats['b10']) > 0
            st.session_state[f"t_6.1_10_meetings"] = is_b10
            if is_b10:
                head = f"проведены рабочие совещания {get_supplier_header(len(b_stats['b10']), 'with')}:"
                st.session_state[f"tx_6.1_10_meetings_{report_id}"] = head + "\n" + "\n".join([f"– {n};" for n in b_stats['b10']])
            
            # Блок 11
            is_b11 = len(b_stats['b11']) > 0
            st.session_state[f"t_6.1_11_approval_sent"] = is_b11
            if is_b11:
                head = f"подготовлены и направлены на согласование проекты Соглашений и протоколов информационного взаимодействия {get_supplier_header(len(b_stats['b11']), 'for')}:"
                st.session_state[f"tx_6.1_11_approval_sent_{report_id}"] = head + "\n" + "\n".join([f"– {n};" for n in b_stats['b11']])

            # Блок 12
            is_b12 = len(b_stats['b12']) > 0
            st.session_state[f"t_6.1_12_agreement_signed"] = is_b12
            if is_b12:
                head = f"подписано Соглашение о доступе, обмене и использовании наборов пространственных данных и протоколы информационного взаимодействия {get_supplier_header(len(b_stats['b12']), 'with')}:"
                st.session_state[f"tx_6.1_12_agreement_signed_{report_id}"] = head + "\n" + "\n".join([f"– {n};" for n in b_stats['b12']])

            # Блок 13 (Протоколы без Соглашения)
            is_b13 = not b_stats['b13'].empty
            st.session_state[f"t_6.1_13_protocol_signed"] = is_b13
            if is_b13:
                lines = []
                for sup, group in b_stats['b13'].groupby('supplier_name'):
                    infos = group['info_name'].tolist()
                    word = "набора" if len(infos) == 1 else "наборов"
                    names = f"”{infos[0]}”" if len(infos) == 1 else ", ".join([f"”{i}”" for i in infos])
                    lines.append(f"– {sup}, {word} {names};")
                head = f"подписан протокол информационного взаимодействия {get_supplier_header(len(lines), 'with')} о размещении на Национальном геопортале набора пространственных данных:"
                st.session_state[f"tx_6.1_13_protocol_signed_{report_id}"] = head + "\n" + "\n".join(lines)

            # Блок 14
            is_b14 = len(b_stats['b14']) > 0
            st.session_state[f"t_6.1_14_letters_gov"] = is_b14
            if is_b14:
                st.session_state[f"tx_6.1_14_letters_gov_{report_id}"] = "направлены письма в Государственный комитет по имуществу по вопросу организации рабочих встреч с поставщиками обязательных наборов пространственных данных, по линиям:\n" + "\n".join([f"– {n};" for n in b_stats['b14']])

            # Блок 15
            is_b15 = len(b_stats['b15']) > 0
            st.session_state[f"t_6.1_15_letters_onpd"] = is_b15
            if is_b15:
                st.session_state[f"tx_6.1_15_letters_onpd_{report_id}"] = f"направлены письма {get_supplier_header(len(b_stats['b15']), 'to')} по вопросу организации рабочих совещаний по вопросам обмена информацией и форматах передачи данных:\n" + "\n".join([f"– {n};" for n in b_stats['b15']])

            # Блок 16
            is_b16 = len(b_stats['b16']) > 0
            st.session_state[f"t_6.1_16_letters_other"] = is_b16
            if is_b16:
                st.session_state[f"tx_6.1_16_letters_other_{report_id}"] = f"направлены письма {get_supplier_header(len(b_stats['b16']), 'for_general')} по теме интеграции данных в Национальную инфраструктуру пространственных данных:\n" + "\n".join([f"– {n};" for n in b_stats['b16']])

    # --- 5. РЕНДЕР КОНСТРУКТОРА (ЦИКЛ ПО ГРУППАМ) ---
    updated_full_data = {}
    
    # 🟢 Загружаем метрики один раз перед циклом, чтобы использовать в полях
    current_metrics = load_report_metrics(report_id)

    for g_key in sorted(sections.keys()):
        g_data = sections[g_key]
        with st.container(border=True):
            st.markdown(f"#### Группа {g_key}")
            
            updated_blocks = {}
            for b_key in sorted(g_data['blocks'].keys()):
                b_data = g_data['blocks'][b_key]
                t_key = f"t_{g_key}_{b_key}"
                tx_key = f"tx_{g_key}_{b_key}_{report_id}"
                
                with st.container(border=False):
                    c_tit, c_tog = st.columns([0.8, 0.2])
                    c_tit.markdown(f"**{b_data['title']}**")
                    is_act = c_tog.toggle("Вкл.", key=t_key, disabled=is_fixed)

                    # 🟢 ПРОВЕРКА: СТРУКТУРИРОВАННЫЙ БЛОК ИЛИ ОБЫЧНЫЙ
                    if b_data.get('is_structured') and not is_fixed:
                        # --- Рендер специализированных полей ввода ---
                        if b_key == "19_stats_631":
                            with st.expander("📈 Ввод метрик посещаемости (6.3.1)", expanded=True):
                                c1, c2, c3 = st.columns(3)
                                c1.number_input("Сайт: Визиты", min_value=0, step=1, key=f"m_{report_id}_s_vis")
                                c2.number_input("Сайт: Уники", min_value=0, step=1, key=f"m_{report_id}_s_unq")
                                c3.number_input("Мета: Визиты", min_value=0, step=1, key=f"m_{report_id}_m_vis")
                                
                                with st.container():
                                    st.caption("Топ-5 разделов сайта")
                                    for i in range(1, 6):
                                        tc1, tc2 = st.columns([0.7, 0.3])
                                        tc1.text_input(f"Раздел {i}", key=f"m_{report_id}_s_p{i}_n")
                                        tc2.number_input(f"Просм. {i}", key=f"m_{report_id}_s_p{i}_v")
                                
                                with st.container():
                                    st.caption("Каталог геоданных и метаданных")
                                    mc1, mc2 = st.columns(2)
                                    mc1.number_input("Мета: Поиск (запросы)", key=f"m_{report_id}_m_sch")
                                    mc2.number_input("Мета: Показы записей", key=f"m_{report_id}_m_rec")
                                    
                                    gc1, gc2, gc3 = st.columns(3)
                                    gc1.number_input("Гео: Визиты", key=f"m_{report_id}_g_vis")
                                    gc2.number_input("Гео: Гости", key=f"m_{report_id}_g_gst")
                                    gc3.number_input("Гео: Рег-ные", key=f"m_{report_id}_g_reg")
                                    
                                    gc4, gc5 = st.columns(2)
                                    gc4.number_input("Гео: Кол-во юзеров", key=f"m_{report_id}_g_u_cnt")
                                    gc5.number_input("Гео: Просмотры карт", key=f"m_{report_id}_g_maps")

                        elif b_key == "20_popularity_632":
                            with st.expander("📊 Ввод популярных наборов (6.3.2)", expanded=True):
                                m_col, g_col = st.columns(2)
                                with m_col:
                                    st.caption("В Каталоге метаданных")
                                    for i in range(1, 6):
                                        st.text_input(f"Мета-набор {i}", key=f"m_{report_id}_pop_m{i}_n")
                                        st.number_input(f"Уник. просм. {i}", key=f"m_{report_id}_pop_m{i}_v")
                                with g_col:
                                    st.caption("В Каталоге геоданных")
                                    for i in range(1, 6):
                                        st.text_input(f"Гео-набор {i}", key=f"m_{report_id}_pop_g{i}_n")
                                        st.number_input(f"Просм. {i}", key=f"m_{report_id}_pop_g{i}_v")

                        # Кнопка генерации текста
                        if st.button("🔄 Сгенерировать текст из цифр", key=f"btn_stich_{b_key}"):
                            # Извлекаем все ключи метрик для этого отчета из session_state
                            prefix = f"m_{report_id}_"
                            local_m = {k.replace(prefix, ''): st.session_state[k] for k in st.session_state if k.startswith(prefix)}
                            
                            # Вызываем функцию склейки (stitch)
                            new_text = stitch_631_text(local_m) if b_key == "19_stats_631" else stitch_632_text(local_m)
                            st.session_state[tx_key] = new_text
                            st.rerun()

                    # Основное текстовое поле (оно есть всегда, но в структурированных блоках оно показывает результат)
                    content = st.text_area("Текст для отчета:", key=tx_key, disabled=not is_act or is_fixed, height=200)
                    
                    updated_blocks[b_key] = {"title": b_data['title'], "active": is_act, "content": content, "is_structured": b_data.get('is_structured', False)}
            
            updated_full_data[g_key] = {"title": g_data['title'], "blocks": updated_blocks}

    # --- 6. КНОПКИ СОХРАНЕНИЯ И СКАЧИВАНИЯ ---
    st.divider()
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if not is_fixed:
            if st.button("💾 Сохранить черновик", width='stretch', type="primary"):
                try:
                    # 1. Сохраняем текст
                    session.execute(
                        text("UPDATE reports_monthly SET sections_data=:d, updated_at=NOW() WHERE report_id=:id"),
                        {"d": json.dumps(updated_full_data, ensure_ascii=False), "id": report_id}
                    )
                    
                    # 2. Сохраняем метрики
                    prefix = f"m_{report_id}_"
                    saved_count = 0
                    for k in st.session_state:
                        if k.startswith(prefix):
                            m_key = k.replace(prefix, '')
                            val = st.session_state[k]
                            if isinstance(val, (int, float)):
                                save_report_metric(session, report_id, m_key, value=val)
                            else:
                                save_report_metric(session, report_id, m_key, text_val=val)
                            saved_count += 1
                    
                    session.commit()
                    clear_cache()
                    # 🟢 Теперь это будет висеть на экране долго
                    st.success(f"✅ Сохранено! Обновлено метрик: {saved_count}")
                except Exception as e:
                    st.error(f"❌ Ошибка сохранения: {e}")
                    session.rollback()
    
    with col2:
        if not is_fixed:
            if st.button("🔒 Зафиксировать", width='stretch', help="Отключает редактирование навсегда"):
                confirm_fix_dialog(session, report_id, updated_full_data)
                
    with col3:
        # Генерация файла
        buf = generate_docx_file(report_date, updated_full_data, MONTHS_RU[s_month])
        st.download_button(
            label="📥 Скачать .docx",
            data=buf,
            file_name=f"Report_NIPD_{s_year}_{s_month}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            width='stretch'
        )

def fetch_provision_stats(start_t, end_t):
    """Сбор статистики с расширенными данными для НИПД и ГКГФ"""
    
    SUCCESS = ('REQ_AGREE_COMPL', 'REQ_CLOSE')
    # Для ГКГФ REQ_CLOSE часто может быть и успехом, и отказом. 
    # Поэтому мы будем смотреть на финальный статус в истории.

    def get_details(req_type):
        # 1. Завершенные (теперь тянем gkf_material_ids)
        query = f"""
            SELECT 
                pr.req_id, pr.applicant_name, pr.request_type, pr.gkf_material_ids,
                it.info_name, ds.dataset_name, sup.supplier_name,
                s.stage_code, h.comments as last_comment,
                (SELECT s2.stage_code FROM provision_request_history h2 
                 JOIN stages s2 ON h2.stage_id = s2.stage_id 
                 WHERE h2.req_id = pr.req_id AND h2.actual_end = h.actual_start 
                 LIMIT 1) as prev_stage_code
            FROM provision_requests pr
            JOIN stages s ON pr.status_id = s.stage_id
            JOIN provision_request_history h ON h.req_id = pr.req_id AND h.stage_id = pr.status_id
            LEFT JOIN info_types it ON pr.nipd_info_id = it.info_id
            LEFT JOIN datasets ds ON it.dataset_id = ds.dataset_id
            LEFT JOIN project_items pi ON it.info_id = pi.info_id
            LEFT JOIN projects p ON pi.project_id = p.project_id
            LEFT JOIN suppliers sup ON p.supplier_id = sup.supplier_id
            WHERE pr.request_type = :rt AND h.actual_start BETWEEN :s AND :e
              AND s.stage_code IN ('REQ_CLOSE', 'REQ_AGREE_COMPL', 'REQ_REGIS_RETUR', 'REQ_REFUS_SUBMI', 'REQ_REFUS_RECEI')
        """
        processed = query_db(query, {"rt": req_type, "s": start_t, "e": end_t})

        # 2. В работе (тянем gkf_material_ids)
        query_work = f"""
            SELECT pr.applicant_name, pr.gkf_material_ids, it.info_name, ds.dataset_name, sup.supplier_name
            FROM provision_requests pr
            JOIN stages s ON pr.status_id = s.stage_id
            LEFT JOIN info_types it ON pr.nipd_info_id = it.info_id
            LEFT JOIN datasets ds ON it.dataset_id = ds.dataset_id
            LEFT JOIN project_items pi ON it.info_id = pi.info_id
            LEFT JOIN projects p ON pi.project_id = p.project_id
            LEFT JOIN suppliers sup ON p.supplier_id = sup.supplier_id
            WHERE pr.request_type = :rt 
              AND s.stage_code NOT IN ('REQ_CLOSE', 'REQ_AGREE_COMPL', 'REQ_REGIS_RETUR', 'REQ_REFUS_SUBMI', 'REQ_REFUS_RECEI')
        """
        in_work = query_db(query_work, {"rt": req_type})

        # 3. Накопительный
        totals = query_db(f"""
            SELECT COUNT(*) as total_received
            FROM provision_requests pr
            WHERE pr.request_type = :rt
        """, {"rt": req_type}).iloc[0]

        return {"processed": processed, "in_work": in_work, "totals": totals}

    return {"NIPD": get_details("НИПД"), "GKF": get_details("Госкартгеофонд")}

def fetch_bureaucracy_logic_stats(start_t, end_t):
    """Сбор статистики для подразделов 6.1.2 (блоки 10-16)"""
    
    def get_sups_by_stage(stage_code, micro_status=4, only_first_iter=False, extra_where=""):
        where_clause = f"stg.stage_code = '{stage_code}'"
        if micro_status:
            where_clause += f" AND ps.micro_status = {micro_status}"
        if only_first_iter:
            where_clause += " AND ps.iteration_count = 1"
        if extra_where:
            where_clause += f" AND {extra_where}"
            
        query = f"""
            SELECT DISTINCT s.supplier_name
            FROM project_stages ps
            JOIN projects p ON ps.project_id = p.project_id
            JOIN suppliers s ON p.supplier_id = s.supplier_id
            JOIN stages stg ON ps.stage_id = stg.stage_id
            WHERE {where_clause} AND ps.actual_end BETWEEN :s AND :e
            ORDER BY s.supplier_name
        """
        return query_db(query, {"s": start_t, "e": end_t})['supplier_name'].tolist()

    # Блок 10: Совещания (Любая итерация, по факту завершения)
    b10_sups = get_sups_by_stage('NEGOTIATIONS')
    
    # Блок 11: Согласование (ТОЛЬКО 1-я итерация, по actual_end)
    b11_sups = get_sups_by_stage('DOCUMENT_APPROVAL', only_first_iter=True, extra_where="s.is_mandatory = TRUE")

    # Блок 12: Подписано Соглашение (is_agreement_project = TRUE)
    b12_sups = get_sups_by_stage('CONTRACT_SIGNED', extra_where="p.is_agreement_project = TRUE")

    # Блок 13: Подписан Протокол (is_agreement_project = FALSE)
    b13_data = query_db("""
        SELECT s.supplier_name, it.info_name
        FROM project_stages ps
        JOIN projects p ON ps.project_id = p.project_id
        JOIN suppliers s ON p.supplier_id = s.supplier_id
        JOIN stages stg ON ps.stage_id = stg.stage_id
        JOIN project_items pi ON p.project_id = pi.project_id
        JOIN info_types it ON pi.info_id = it.info_id
        WHERE stg.stage_code = 'CONTRACT_SIGNED' 
          AND ps.micro_status = 4 
          AND p.is_agreement_project = FALSE
          AND ps.actual_end BETWEEN :s AND :e
    """, {"s": start_t, "e": end_t})
    
    # Блоки 14, 15, 16: Запрос переговоров (ТОЛЬКО 1-я итерация, по actual_end)
    b14_sups = get_sups_by_stage('NEGOTIATIONS_REQUEST', only_first_iter=True, extra_where="s.is_gov_agency = TRUE AND s.is_mandatory = TRUE")
    b15_sups = get_sups_by_stage('NEGOTIATIONS_REQUEST', only_first_iter=True, extra_where="s.is_gov_agency = FALSE AND s.is_mandatory = TRUE")
    b16_sups = get_sups_by_stage('NEGOTIATIONS_REQUEST', only_first_iter=True, extra_where="s.is_mandatory = FALSE")

    return {
        "b10": b10_sups, "b11": b11_sups, "b12": b12_sups, 
        "b13": b13_data, "b14": b14_sups, "b15": b15_sups, "b16": b16_sups
    }

def stitch_631_text(m):
    """Собирает текст для блока 6.3.1"""
    # Сайт
    site = [
        f"Сайт Национального геопортала nipd.by:",
        f"– {int(m.get('s_vis', 0))} посещений;",
        f"– {int(m.get('s_unq', 0))} уникальных посетителей;",
        f"– наиболее популярные разделы и страницы сайта:"
    ]
    for i in range(1, 6):
        name = m.get(f's_p{i}_n')
        val = m.get(f's_p{i}_v')
        if name and val: site.append(f"\t– раздел ”{name}” – {int(val)} {pluralize(val, ['просмотр', 'просмотра', 'просмотров'])};")
    
    # Каталог метаданных
    meta = [
        f"Каталог метаданных:",
        f"– {int(m.get('m_vis', 0))} посещений;",
        f"– {int(m.get('m_sch', 0))} поисковых запроса;",
        f"– {int(m.get('m_rec', 0))} уникальных показа записей метаданных;"
    ]
    
    # Каталог геоданных
    geo = [
        f"Каталог геоданных:",
        f"– {int(m.get('g_vis', 0))} посещений, в том числе:",
        f"\t– {int(m.get('g_gst', 0))} – незарегистрированными,",
        f"\t– {int(m.get('g_reg', 0))} – зарегистрированными пользователями (всего {int(m.get('g_u_cnt', 0))} {pluralize(m.get('g_u_cnt', 0), ['пользователь', 'пользователя', 'пользователей'])});",
        f"– {int(m.get('g_maps', 0))} {pluralize(m.get('g_maps', 0), ['просмотр', 'просмотра', 'просмотров'])} карт."
    ]
    
    header = "6.3.1. Анализ посещаемости и вовлеченности посетителей Национального геопортала – осуществлялся сбор и сохранение сведений о статистике посещаемости компонентов Национального геопортала. За отчетный период зафиксированы следующие показатели:"
    return header + "\n" + "\n".join(site + meta + geo)

def stitch_632_text(m):
    """Собирает текст для блока 6.3.2"""
    header = "6.3.2. Оценка востребованности наборов пространственных данных, включенных в НИПД – осуществлялся сбор и сохранение сведений о статистике просмотров наборов пространственных данных, включенных в НИПД. Была зафиксирована наиболее востребованная информация (количество просмотров):"
    
    lines = ["в Каталоге метаданных:"]
    for i in range(1, 6):
        n, v = m.get(f'pop_m{i}_n'), m.get(f'pop_m{i}_v')
        if n: lines.append(f"– {n} – {int(v or 0)} уникальных {pluralize(v or 0, ['просмотр', 'просмотра', 'просмотров'])};")
    
    lines.append("в Каталоге геоданных:")
    for i in range(1, 6):
        n, v = m.get(f'pop_g{i}_n'), m.get(f'pop_g{i}_v')
        if n: lines.append(f"– {n} – {int(v or 0)} {pluralize(v or 0, ['просмотр', 'просмотра', 'просмотров'])};")
        
    return header + "\n" + "\n".join(lines)